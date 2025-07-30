import os
import requests
import pandas as pd
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv(override=True)

def get_greeting():
    """Returns a greeting based on the time of day."""
    hour = datetime.now().hour
    if hour < 12:
        return "Good Morning"
    elif hour < 17:
        return "Good Afternoon"
    else:
        return "Good Evening"

def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            element = OxmlElement(f'w:{edge}')
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def read_timesheet_data(filename: str) -> str:
    """
    Use this to read invoice data from an Excel timesheet file (XLSX format). 
    It returns rows of date, status, and remarks. Input must be the filename (e.g., 'timesheet_july.xlsx').
    """
    try:
        # Resolve the full path
        filepath = os.path.join(os.getcwd(), filename)
        
        # If file does not exist, create an empty timesheet
        if not os.path.exists(filepath):
            # Define initial structure
            df = pd.DataFrame(columns=["Date", "Status", "Remarks"])
            df.to_excel(filepath, index=False)

        # Read the Excel file
        df = pd.read_excel(filepath)

        if df.empty:
            return "The Excel file is empty."

        # Convert the DataFrame into a nicely formatted table-like string
        lines = []
        for _, row in df.iterrows():
            date = row.get("Date", "")
            status = row.get("Status", "")
            remarks = row.get("Remarks", "")
            lines.append(f"{date} | {status} | {remarks}")

        return "Timesheet Records:\n" + "\n".join(lines)

    except Exception as e:
        return f"Error reading Excel timesheet: {str(e)}"

def create_invoice_document(filename: str, data: dict) -> str:
    """    
    Generates a .docx invoice from a dictionary of data.    
    Use this to generate and save a formatted invoice as a Word (.docx) file. 
    Input must be a dictionary with 'filename' and 'data' keys.
    """
    try:

        if not filename or not data:
            return "Error: Both 'filename' and 'data' are required."

        if not filename.endswith(".docx"):
            filename += ".docx"

        # Use an absolute path or a designated output directory
        output_dir = os.path.join(os.getcwd())
        file_path = os.path.join(output_dir, filename)
        
        doc = Document()

        # --- Set Default Font Style ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        # --- Main Table (1 column, used for layout) ---
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(6.5)

        # --- Row 1: INVOICE Title ---
        cell = table.cell(0, 0)
        p = cell.paragraphs[0]
        run = p.add_run("INVOICE")
        run.font.name = "Arial Black"
        run.font.size = Pt(28)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"}, top={"sz": 6, "val": "single"})

        # --- Row 2: Name (left) and Date (right) ---
        row_cell = table.add_row().cells[0]
        inner_table = row_cell.add_table(rows=1, cols=2)
        inner_table.columns[0].width = Inches(4.0)
        inner_table.columns[1].width = Inches(2.5)
        
        left_cell = inner_table.cell(0, 0)
        left_cell.paragraphs[0].add_run(data["name"]).bold = True
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        right_cell = inner_table.cell(0, 1)
        right_cell.paragraphs[0].add_run(f"{data['date']}")
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in inner_table._cells:
            set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
        set_cell_border(row_cell, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"}, top={"val": "nil"})

        # --- Row 3: Bill To Block ---
        cell = table.add_row().cells[0]
        p = cell.paragraphs[0]
        p.add_run("Bill To:\n").bold = True
        for line in data["bill_to"]:
            p.add_run(f"    {line}\n")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # --- Row 4: Details Table Headers ---
        row_cell = table.add_row().cells[0]
        header_table = row_cell.add_table(rows=1, cols=2)
        header_table.columns[0].width = Inches(5.0)
        header_table.columns[1].width = Inches(1.5)
        desc_cell = header_table.cell(0, 0)
        amt_cell = header_table.cell(0, 1)
        header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        header_table.autofit = False

        desc_cell.text = "DESCRIPTION"
        amt_cell.text = "AMOUNT"

        for cell in [desc_cell, amt_cell]:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), "ff99cc")
            cell._tc.get_or_add_tcPr().append(shading_elm)
        set_cell_border(row_cell, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # --- Row 5: Details Table Content ---
        row_cell = table.add_row().cells[0]
        details_table = row_cell.add_table(rows=0, cols=2)
        details_table.columns[0].width = Inches(5.0)
        details_table.columns[1].width = Inches(1.5)

        cells = details_table.add_row().cells
        cells[0].text = data["salary_description"]
        cells[1].text = ""

        # Add itemized details, handling both dict and string formats for robustness
        for item in data["details"]:
            cells = details_table.add_row().cells
            if isinstance(item, dict):
                # Handle the new, preferred dictionary format
                cells[0].text = item.get("description", "")
                cells[1].text = item.get("amount", "")
            elif isinstance(item, str) and ':' in item:
                # Handle the old string format for backward compatibility
                key, value = item.split(":", 1)
                cells[0].text = key.strip()
                cells[1].text = value.strip()
            else:
                # Handle any other unexpected format
                cells[0].text = str(item)
                cells[1].text = ""
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        set_cell_border(row_cell, top={"val": "nil"}, bottom={"sz": 6, "val": "single"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        # --- Row 6: TOTAL Row ---
        row_cell = table.add_row().cells[0]
        total_table = row_cell.add_table(rows=1, cols=2)
        total_table.columns[0].width = Inches(5.0)
        total_table.columns[1].width = Inches(1.5)
        left = total_table.cell(0, 0)
        right = total_table.cell(0, 1)

        left.paragraphs[0].add_run("TOTAL").bold = True
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        right.paragraphs[0].add_run(data["total"]).bold = True
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "ffcc99")
        right._tc.get_or_add_tcPr().append(shading_elm)
        right.alignment = WD_TABLE_ALIGNMENT.LEFT
        right.autofit = False

        set_cell_border(row_cell, top={"sz": 6, "val": "single"}, bottom={"val": "nil"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})
        set_cell_border(right, left={"sz": 6, "val": "single"}, bottom={"sz": 6, "val": "single"})

        # --- Row 7: Amount in Words ---
        cell = table.add_row().cells[0]
        p = cell.paragraphs[0]
        p.add_run("Amount in Words: ").bold = True
        p.add_run(data["total_words"])
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_border(cell, top={"val": "nil"}, bottom={"sz": 6, "val": "single"}, left={"sz": 6, "val": "single"}, right={"sz": 6, "val": "single"})

        doc.save(file_path)
        return f"Invoice successfully written to {file_path}"

    except Exception as e:
        return f"Failed to write Word document: {str(e)}"

def save_or_update_timesheet(filename: str, date: str, status: str, remarks: str) -> str:
    """
    Saves or updates a single entry in the Excel timesheet.
    If an entry for the given date already exists, it will be UPDATED.
    Otherwise, a new entry will be ADDED.
    The date must be in 'YYYY-MM-DD' format.
    It needs the filename, date, status, and remarks as input. Filename usually be in the format: 'timesheet_<month>.xlsx').
    """
    try:
        filepath = os.path.join(os.getcwd(), filename)
        
        # Define the columns to ensure consistency
        columns = ["Date", "Status", "Remarks"]

        # Check if the file exists and read it; otherwise, create an empty DataFrame
        if os.path.exists(filepath):
            df = pd.read_excel(filepath)
            # Ensure the Date column is treated as string to avoid formatting issues
            df['Date'] = df['Date'].astype(str).str.split(' ').str[0]
        else:
            df = pd.DataFrame(columns=columns)

        # Check if an entry for the given date already exists
        if date in df['Date'].values:
            # Update existing entry
            idx = df.index[df['Date'] == date][0]
            df.loc[idx, 'Status'] = status
            df.loc[idx, 'Remarks'] = remarks
            action = "updated"
        else:
            # Add new entry
            new_entry = pd.DataFrame([{"Date": date, "Status": status, "Remarks": remarks}])
            df = pd.concat([df, new_entry], ignore_index=True)
            action = "added"
            
        # Save the updated DataFrame back to Excel
        df.to_excel(filepath, index=False)
        
        return f"Success: The entry for {date} was {action} in {filename}."

    except Exception as e:
        return f"Error modifying Excel timesheet: {str(e)}"

def send_message_with_attachments(
    xlsx_filename: str, 
    docx_filename: str
) -> str:
    """
    Sends an email with an XLSX and a DOCX file as attachments from the current directory.
    
    Args:
        xlsx_filename: The filename of the .xlsx file (e.g., 'timesheet_july.xlsx').
        docx_filename: The filename of the .docx file (e.g., 'invoice_july.docx').
        
    Returns:
        A string indicating success or failure.
    """
    try:
        telegram_token = os.getenv("TELEGRAM_BOT_TOKEN")
        telegram_chat_id = os.getenv("TELEGRAM_CHAT_ID")

        if not all([telegram_token, telegram_chat_id]):
            raise ValueError("Missing TELEGRAM_BOT_TOKEN or TELEGRAM_CHAT_ID environment variables")

        greeting = get_greeting()
        message_text = (
            f"Hi,\n{greeting}.\n\n"
            "I've attached the timesheet and invoice for the month of July.\n"
            "Please review and approve at your convenience."
        )

        send_message_url = f"https://api.telegram.org/bot{telegram_token}/sendMessage"
        requests.post(send_message_url, data={
            "chat_id": telegram_chat_id,
            "text": message_text
        })

        files_to_send = [xlsx_filename, docx_filename]
        sent_count = 0

        for file in files_to_send:
            if os.path.exists(file):
                with open(file, 'rb') as f:
                    send_file_url = f"https://api.telegram.org/bot{telegram_token}/sendDocument"
                    response = requests.post(send_file_url, data={
                        "chat_id": telegram_chat_id
                    }, files={
                        "document": (file, f)
                    })
                    if response.status_code == 200:
                        sent_count += 1
                    else:
                        return f"Failure: Telegram API error for file {file}: {response.text}"
        
        if sent_count == 0:
            return "Failure: No valid files found to send via Telegram."

        return f"Success: Sent {sent_count} file(s) to Telegram chat."

    except Exception as e:
        return f"An error occurred: {e}"